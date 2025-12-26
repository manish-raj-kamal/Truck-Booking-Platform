from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_heading(doc, text, level):
    h = doc.add_heading(text, level)
    run = h.runs[0]
    run.font.color.rgb = RGBColor(0, 0, 0)  # Black color
    return h

def add_monospaced_paragraph(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    # Force Courier New for compatibility
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), 'Courier New')

def create_srs_doc():
    doc = Document()
    
    # Title Page
    title = doc.add_heading('Software Requirements Specification', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('TruckSuvidha - Truck Booking System')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(16)
    
    doc.add_page_break()

    # Revision History
    add_heading(doc, 'Revision History', 1)
    table = doc.add_table(rows=2, cols=4)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Version'
    hdr_cells[1].text = 'Date'
    hdr_cells[2].text = 'Author'
    hdr_cells[3].text = 'Description'
    
    row_cells = table.rows[1].cells
    row_cells[0].text = '1.0'
    row_cells[1].text = 'December 25, 2024'
    row_cells[2].text = 'Development Team'
    row_cells[3].text = 'Initial SRS Document'
    
    doc.add_paragraph()

    # Table of Contents Placeholder
    add_heading(doc, 'Table of Contents', 1)
    doc.add_paragraph('(Please update the Table of Contents in Word automatically via References > Update Table)')
    doc.add_page_break()

    # 1. Introduction
    add_heading(doc, '1. Introduction', 1)
    doc.add_paragraph('The TruckSuvidha platform is a comprehensive web-based logistics solution designed to connect shippers with transporters across India. This SRS document provides a detailed description of the system requirements.')

    # 1.1 Purpose
    add_heading(doc, '1.1 Purpose', 2)
    doc.add_paragraph('The purpose of this SRS is to define the requirements for the TruckSuvidha platform.')
    doc.add_paragraph('Product Name: TruckSuvidha - Truck Booking Platform', style='List Bullet')
    
    doc.add_paragraph('What the software will do:')
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Enable customers to post loads/shipments with detailed requirements.')
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Allow transporters/drivers to browse and bid on available loads.')
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Facilitate secure payment processing through Razorpay integration.')
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Provide real-time order status tracking throughout the delivery lifecycle.')
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Support multiple authentication methods including email/password and Google OAuth.')

    doc.add_paragraph('What the software will NOT do:')
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Provide physical transportation services directly.')
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Handle insurance claims or disputes.')
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Provide real-time GPS vehicle tracking (status-based tracking only).')

    # 1.2 Scope
    add_heading(doc, '1.2 Scope', 2)
    doc.add_paragraph('Product Overview: TruckSuvidha is a full-stack web application serving as a digital marketplace for logistics services in India.')
    doc.add_paragraph('System Boundaries:')
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Web-based platform accessible via modern browsers.')
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Backend deployed on Vercel.')
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Database hosted on MongoDB Atlas.')
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Payments handled through Razorpay.')

    # 1.3 Definitions
    add_heading(doc, '1.3 Definitions, Acronyms, and Abbreviations', 2)
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Term'
    hdr_cells[1].text = 'Definition'
    
    definitions = [
        ('SRS', 'Software Requirements Specification'),
        ('API', 'Application Programming Interface'),
        ('JWT', 'JSON Web Token - Used for secure authentication'),
        ('OAuth', 'Open Authorization - Standard for token-based authentication'),
        ('REST', 'Representational State Transfer - API architecture style'),
        ('MT', 'Metric Ton - Unit of weight measurement'),
        ('OTP', 'One-Time Password'),
        ('Load', 'Shipment or cargo that needs to be transported'),
        ('Quote', 'Price bid submitted by a transporter for a load'),
        ('Transporter', 'Driver or trucking company providing transport services'),
        ('Shipper', 'Customer who needs goods transported')
    ]
    
    for term, definition in definitions:
        row_cells = table.add_row().cells
        row_cells[0].text = term
        row_cells[1].text = definition

    # 1.4 References
    add_heading(doc, '1.4 References', 2)
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Document/Resource'
    hdr_cells[1].text = 'Description'
    
    refs = [
        ('MongoDB Documentation', 'https://docs.mongodb.com/'),
        ('Express.js Documentation', 'https://expressjs.com/'),
        ('React Documentation', 'https://react.dev/'),
        ('Razorpay API Documentation', 'https://razorpay.com/docs/'),
        ('JWT.io', 'https://jwt.io/')
    ]
    for ref, link in refs:
        row_cells = table.add_row().cells
        row_cells[0].text = ref
        row_cells[1].text = link

    # 1.5 Overview
    add_heading(doc, '1.5 Overview', 2)
    doc.add_paragraph('This SRS is organized into sections covering the general description, specific functional/non-functional requirements, analysis models, and administrative details.')

    # 2. General Description
    add_heading(doc, '2. General Description', 1)

    # 2.1 Product Perspective
    add_heading(doc, '2.1 Product Perspective', 2)
    doc.add_paragraph('TruckSuvidha is a standalone web-based logistics platform.')
    p = doc.add_paragraph()
    run = p.add_run('System Architecture:')
    run.bold = True
    
    arch_diagram = """
┌───────────────────────────────────────┐
│              CLIENT LAYER             │
│   React 18 + Vite Frontend (SPA)      │
└───────────────────────────────────────┘
                   │
                   ▼
┌───────────────────────────────────────┐
│              SERVER LAYER             │
│ Node.js + Express.js Backend (REST)   │
└───────────────────────────────────────┘
                   │
    ┌──────────────┼──────────────┐
    ▼              ▼              ▼
┌─────────┐   ┌─────────┐   ┌─────────┐
│ MongoDB │   │ Razorpay│   │ Google  │
│  Atlas  │   │         │   │  OAuth  │
└─────────┘   └─────────┘   └─────────┘
"""
    add_monospaced_paragraph(doc, arch_diagram)

    # 2.2 Product Functions
    add_heading(doc, '2.2 Product Functions', 2)
    funcs = [
        "User Management: Registration, auth, profile mgmt, RBAC.",
        "Load Management: Posting, viewing, lifecycle tracking.",
        "Quote/Bidding: Driver bidding, acceptance, assignment.",
        "Order Lifecycle: Status progression (Open -> Delivered).",
        "Payments: Dynamic fee calculation, Razorpay integration.",
        "Admin Functions: Manage users, loads, trucks, content."
    ]
    for f in funcs:
        doc.add_paragraph(f, style='List Bullet')

    # 2.3 User Characteristics
    add_heading(doc, '2.3 User Characteristics', 2)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'User Type'
    hdr_cells[1].text = 'Expertise'
    hdr_cells[2].text = 'Primary Tasks'
    
    users = [
        ('Customer', 'Basic/Intermediate', 'Post loads, make payments, track orders'),
        ('Driver', 'Basic (Mobile)', 'Browse loads, submit quotes, update status'),
        ('Admin', 'Intermediate', 'Manage users, loads, content'),
        ('SuperAdmin', 'Advanced', 'Full system access')
    ]
    for u in users:
        row = table.add_row().cells
        row[0].text = u[0]
        row[1].text = u[1]
        row[2].text = u[2]

    # 2.4 General Constraints
    add_heading(doc, '2.4 General Constraints', 2)
    doc.add_paragraph('Hardware: Device with modern browser/internet.', style='List Bullet')
    doc.add_paragraph('Software: Chrome v90+, Firefox v88+, Safari v14+.', style='List Bullet')
    doc.add_paragraph('Security: HTTPS, bcrypt, JWT 24h expiry.', style='List Bullet')
    doc.add_paragraph('Regulatory: GST and Indian data protection compliance.', style='List Bullet')

    # 2.5 Assumptions
    add_heading(doc, '2.5 Assumptions and Dependencies', 2)
    doc.add_paragraph('Assumptions: Stable internet, valid emails, INR currency.', style='List Bullet')
    doc.add_paragraph('Dependencies: MongoDB Atlas, Razorpay, Google OAuth, Vercel.', style='List Bullet')

    # 3. Specific Requirements
    add_heading(doc, '3. Specific Requirements', 1)

    # 3.1 External Interface
    add_heading(doc, '3.1 External Interface Requirements', 2)
    doc.add_paragraph('3.1.1 User Interfaces: Responsive (320px-2560px). Pages: Landing, Login, Load Board, Post Load, Profile, Admin.', style='List Paragraph')
    doc.add_paragraph('3.1.2 Hardware Interfaces: Camera access (avatar), Touch input support.', style='List Paragraph')
    doc.add_paragraph('3.1.3 Software Interfaces: MongoDB Wire Protocol, Razorpay API, Google OAuth.', style='List Paragraph')
    doc.add_paragraph('3.1.4 Communications: REST over HTTPS (TLS 1.2), JSON format, SMTP (Email).', style='List Paragraph')

    # 3.2 Functional Requirements
    add_heading(doc, '3.2 Functional Requirements', 2)

    # 3.2.1 Registration
    add_heading(doc, '3.2.1 User Registration and Authentication', 3)
    doc.add_paragraph('Inputs: Email, Password, Name, Role (Customer/Driver), Google OAuth credentials.')
    doc.add_paragraph('Processing: Validate email, hash password (bcrypt), generate JWT.')
    doc.add_paragraph('Error Handling: 400 (Invalid), 409 (Email exists), 401 (Credentials).')

    # 3.2.2 Load Posting
    add_heading(doc, '3.2.2 Load Posting', 3)
    doc.add_paragraph('Inputs: Load Type, Source/Dest, Material, Weight, Truck Type, Date.')
    doc.add_paragraph('Processing: Calculate fee (Base+Weight+Material), Create Razorpay Order, Verify Signature.')
    doc.add_paragraph('Outputs: Load Object, Payment Record.')

    # 3.2.3 Quote System
    add_heading(doc, '3.2.3 Quote/Bidding System', 3)
    doc.add_paragraph('Inputs: LoadID, Amount, Message, Delivery Days.')
    doc.add_paragraph('Processing: Verify driver role. On accept: assign driver, reject others, update load status.')
    doc.add_paragraph('Error Handling: 403 (Non-driver/Own load), 409 (Duplicate).')

    # 3.2.4 Order Status
    add_heading(doc, '3.2.4 Order Status Management', 3)
    doc.add_paragraph('Status Flow: Open -> Quoted -> Assigned -> Picked Up -> In Transit -> Delivered -> Completed.')
    doc.add_paragraph('Processing: Validate permission and status transition logic.')

    # 3.2.5 Payments
    add_heading(doc, '3.2.5 Payment Processing', 3)
    doc.add_paragraph('Formula: min(Base + WeightFee + MaterialFee + TruckFee, 1000).')
    doc.add_paragraph('Processing: Verify HMAC SHA256 signature from Razorpay.')

    # 3.5 Non-Functional
    add_heading(doc, '3.5 Non-Functional Requirements', 2)
    doc.add_paragraph('3.5.1 Performance: API < 500ms, Page Load < 3s, 100+ Concurrent users.')
    doc.add_paragraph('3.5.2 Reliability: 99% Uptime, Zero data loss.')
    doc.add_paragraph('3.5.3 Availability: 24/7 Service, Max 4h/month maintenance.')
    doc.add_paragraph('3.5.4 Security: TLS 1.2+, JWT Auth, RBAC, CSRF/XSS protection.')
    doc.add_paragraph('3.5.5 Maintainability: Modular MVC, JSDoc.')
    doc.add_paragraph('3.5.6 Portability: Mobile responsive, browser agnostic.')

    # 3.7 Constraints
    add_heading(doc, '3.7 Design Constraints', 2)
    doc.add_paragraph('Stack: React 18, Node.js, MongoDB.')
    doc.add_paragraph('Deployment: Vercel (10s serverless timeout).')
    doc.add_paragraph('Language: English UI, Hindi Truck sizes.')

    # 4 Analysis Models
    add_heading(doc, '4. Analysis Models', 1)
    add_heading(doc, '4.1 Data Flow Diagrams (DFD)', 2)
    
    dfd_lvl0 = """
     ┌──────────────┐                       ┌──────────────┐
     │   Customer   │◄──────┐       ┌──────►│    Driver    │
     │  (Shipper)   │       │       │       │ (Transporter)│
     └──────────────┘       ▼       ▼       └──────────────┘
                      ┌──────────────────┐
                      │   TruckSuvidha   │
                      │     Platform     │
                      └──────────────────┘
                            │      │
          ┌──────────────┐  │      │  ┌──────────────┐
          │   Razorpay   │◄─┘      └─►│    Admin     │
          │  (Payments)  │            │  (Management)│
          └──────────────┘            └──────────────┘
"""
    add_monospaced_paragraph(doc, 'Level 0 - Context Diagram')
    add_monospaced_paragraph(doc, dfd_lvl0)

    # 5-11 Administrative
    doc.add_page_break()
    admin_headers = [
        "5. GitHub Link",
        "6. Deployed Link",
        "7. Client Approval Proof",
        "8. Client Location Proof",
        "9. Transaction ID Proof",
        "10. Email Acknowledgement",
        "11. GST No."
    ]
    
    admin_content = [
        "Repository URL: https://github.com/[username]/Truck-Booking",
        "Live Application URL: https://somya-truck-booking.vercel.app",
        "[TODO: Insert Proof]",
        "[TODO: Insert Proof]",
        "[TODO: Insert Proof]",
        "[TODO: Insert Proof]",
        "GST Number: [TODO]"
    ]

    for h, c in zip(admin_headers, admin_content):
        add_heading(doc, h, 1)
        doc.add_paragraph(c)

    # Appendices
    doc.add_page_break()
    add_heading(doc, 'A. Appendices', 1)
    
    add_heading(doc, 'A.1 Appendix 1 - Technology Stack', 2)
    doc.add_paragraph('Frontend: React 18.3.1, Vite, TailwindCSS, React Router, TanStack Query.')
    doc.add_paragraph('Backend: Node.js, Express, Mongoose, Passport.js, Razorpay.')

    add_heading(doc, 'A.2 Appendix 2 - User Role Permissions', 2)
    doc.add_paragraph('(See detailed matrix in main document requirements)')

    # Save
    doc.save('TruckSuvidha_SRS.docx')
    print("Document created successfully: TruckSuvidha_SRS.docx")

if __name__ == "__main__":
    create_srs_doc()